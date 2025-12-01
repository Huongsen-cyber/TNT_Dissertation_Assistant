import streamlit as st
import traceback

# --- 1. CẤU HÌNH TRANG ---
st.set_page_config(
    page_title="TNT Dissertation Master AI",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================
# 0. BỘ NÃO TNT V1.1 (NHÚNG TRI THỨC)
# ==========================================
TNT_MASTER_PROMPT = """
ROLE: You are "TNT Advanced AI Editor & Writer V1.1", a specialized assistant for Doctoral Dissertations.

[TNT COMMAND SYSTEM - KNOWLEDGE BASE]:
You must map User's Natural Language requests to these specific Command Codes:

1. ANALYSIS (Phân tích & Kiểm tra):
   - WF-DMAI: Deep Structural Analysis (Phân tích cấu trúc, chia đoạn, tìm lỗ hổng).
   - WF-QACHECK: Check logic, flow, coherence. (Dùng khi user hỏi "bài này có lỗi logic không?", "kiểm tra mạch văn").

2. EDITING (Biên tập & Sửa chữa):
   - ED-STD: Standard Academic Editing. (Dùng khi user bảo "sửa lỗi chính tả", "làm văn phong hay hơn").
   - ED-EXT25: Expand analysis (+25%). (Dùng khi user bảo "viết sâu hơn", "mở rộng ý này", "thêm luận cứ").
   - ED-RED05: Condense text. (Dùng khi user bảo "rút gọn", "viết súc tích lại").

3. WRITING (Viết mới):
   - WF-GENDRAFT: Generate new content. (Dùng khi user bảo "viết cho tôi chương này", "soạn thảo mục này").

4. FORMATTING (Định dạng):
   - FMT-FNAF02: Standard 5-page chunk format with Glossary/Footnotes. (Luôn dùng định dạng này cho đầu ra chính thức).

[SMART AGENT BEHAVIOR]:
If the user says: "Hãy sửa lại chương này cho tôi", you reply:
"🔍 **Phân tích:** Bạn muốn chỉnh sửa văn phong và ngữ pháp.
🛠️ **Kích hoạt lệnh:** `ED-STD` + `FMT-FNAF02`
... [Then execute the task] ..."
"""

# --- BẮT ĐẦU KHỐI CODE ---
try:
    import google.generativeai as genai
    from pypdf import PdfReader
    from docx import Document
    from io import BytesIO
    import json
    import os
    import tempfile
    import datetime
    
    from streamlit_mic_recorder import mic_recorder
    from gtts import gTTS
    from pydub import AudioSegment
    from google.oauth2.credentials import Credentials
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
    import speech_recognition as sr

    # --- ID THƯ MỤC GỐC ---
    ROOT_FOLDER_ID = "1eojKKKoMk4uLBCLfCpVhgWnaoTtOiu8p"

    # ==========================================
    # CÁC HÀM XỬ LÝ (DRIVE, FILE)
    # ==========================================
    def get_drive_service():
        if "oauth_token" not in st.secrets:
            st.error("❌ Lỗi: Chưa cấu hình 'oauth_token' trong Secrets!")
            return None
        try:
            token_info = json.loads(st.secrets["oauth_token"])
            creds = Credentials.from_authorized_user_info(token_info)
            return build('drive', 'v3', credentials=creds)
        except Exception as e:
            st.error(f"❌ Lỗi xác thực Google: {e}")
            return None

    def upload_to_drive(file_obj, filename, folder_id):
        try:
            service = get_drive_service()
            if not service: return None, "Lỗi kết nối"
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M")
            final_filename = f"{filename.replace('.docx', '')}_{timestamp}.docx"
            file_metadata = {'name': final_filename, 'parents': [folder_id]}
            file_obj.seek(0)
            media = MediaIoBaseUpload(file_obj, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
            return file.get('id'), final_filename
        except Exception as e: return None, str(e)

    # --- HÀM ĐỆ QUY: LẤY CÂY THƯ MỤC ---
    def get_all_folders_recursive(service, parent_id, prefix=""):
        folders = []
        try:
            results = service.files().list(
                q=f"'{parent_id}' in parents and mimeType = 'application/vnd.google-apps.folder' and trashed=false",
                fields="files(id, name)", orderBy="name").execute()
            for item in results.get('files', []):
                folders.append({'id': item['id'], 'name': f"{prefix}📁 {item['name']}"})
                folders.extend(get_all_folders_recursive(service, item['id'], prefix + "-- "))
        except: pass
        return folders

    # --- HÀM ĐỆ QUY: LẤY FILE (DEEP SCAN) ---
    def list_files_deep(service, folder_id):
        all_files = []
        try:
            # Lấy file thư mục hiện tại
            files = service.files().list(
                q=f"'{folder_id}' in parents and mimeType != 'application/vnd.google-apps.folder' and trashed=false",
                fields="files(id, name, mimeType)", orderBy="name").execute().get('files', [])
            all_files.extend(files)
            
            # Lấy file thư mục con
            subfolders = service.files().list(
                q=f"'{folder_id}' in parents and mimeType = 'application/vnd.google-apps.folder' and trashed=false",
                fields="files(id)"
            ).execute().get('files', [])
            
            for sub in subfolders:
                all_files.extend(list_files_deep(service, sub['id']))
        except: pass
        return all_files

    def read_drive_file(service, file_id, filename, mimeType):
        try:
            file_stream = BytesIO()
            if mimeType == 'application/vnd.google-apps.document':
                request = service.files().export_media(fileId=file_id, mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            else:
                request = service.files().get_media(fileId=file_id)
            
            downloader = MediaIoBaseDownload(file_stream, request)
            done = False
            while done is False: status, done = downloader.next_chunk()
            file_stream.seek(0)
            
            if filename.endswith(".pdf") or mimeType == 'application/pdf': return get_pdf_content(file_stream)
            else: return get_docx_content(file_stream)
        except Exception as e: return "" 

    def get_pdf_content(f):
        try:
            reader = PdfReader(f); text = ""
            for p in reader.pages: text += p.extract_text() + "\n"
            return text
        except: return ""

    def get_docx_content(f):
        try:
            doc = Document(f)
            return "\n".join([p.text for p in doc.paragraphs])
        except: return ""

    def get_local_content(f):
        f.seek(0)
        if f.name.endswith(".pdf"): return get_pdf_content(f)
        else: return get_docx_content(f)

    # ==========================================
    # QUẢN LÝ SESSION STATE (BỘ NHỚ)
    # ==========================================
    if 'global_context' not in st.session_state: st.session_state.global_context = ""
    if 'read_history' not in st.session_state: st.session_state.read_history = [] # Lưu danh sách tên file đã đọc
    if 'current_folder_id' not in st.session_state: st.session_state.current_folder_id = ROOT_FOLDER_ID
    if 'current_folder_name' not in st.session_state: st.session_state.current_folder_name = "Thư mục gốc"
    if 'folder_tree_cache' not in st.session_state: st.session_state.folder_tree_cache = []
    if 'found_files_cache' not in st.session_state: st.session_state.found_files_cache = []

    # ==========================================
    # GIAO DIỆN SIDEBAR
    # ==========================================
    with st.sidebar:
        st.title("🎙️ TNT Smart Center")
        api_key = st.text_input("Nhập Google AI Key:", type="password")
        
        st.divider()
        audio_bytes = mic_recorder(start_prompt="🔴 Ghi âm", stop_prompt="⏹️ Dừng", key='recorder')
        
        st.divider()
        # HIỂN THỊ TRẠNG THÁI BỘ NHỚ THÔNG MINH
        with st.expander(f"🧠 Bộ nhớ: {len(st.session_state.read_history)} file", expanded=False):
            if st.session_state.read_history:
                st.write("**Các file đã nạp:**")
                for f in st.session_state.read_history:
                    st.caption(f"✅ {f}")
                if st.button("🗑️ Quên tất cả (Reset)"):
                    st.session_state.global_context = ""
                    st.session_state.read_history = []
                    st.rerun()
            else:
                st.write("(Chưa có dữ liệu)")

        st.divider()
        st.subheader("📂 Quản lý Dữ liệu")
        source_option = st.radio("Nguồn:", ["Tải từ máy tính", "📁 Duyệt Google Drive"])

        # 1. TẢI TỪ MÁY (CÓ CHECK TRÙNG LẶP)
        if source_option == "Tải từ máy tính":
            uploaded_files = st.file_uploader("Chọn file:", type=["pdf", "docx"], accept_multiple_files=True)
            if uploaded_files:
                if st.button("🚀 Nạp dữ liệu mới"):
                    with st.spinner("Đang xử lý..."):
                        new_ctx = ""
                        count_new = 0
                        count_skip = 0
                        
                        for f in uploaded_files:
                            # --- LOGIC CHỐNG TRÙNG ---
                            if f.name in st.session_state.read_history:
                                count_skip += 1
                                continue # Bỏ qua file này
                            
                            # Nếu chưa có thì xử lý
                            upload_to_drive(f, f.name, ROOT_FOLDER_ID)
                            new_ctx += f"\n=== TÀI LIỆU: {f.name} ===\n{get_local_content(f)}\n"
                            st.session_state.read_history.append(f.name)
                            count_new += 1
                        
                        if count_new > 0:
                            st.session_state.global_context += new_ctx
                            st.success(f"✅ Đã nạp thêm {count_new} file mới.")
                        
                        if count_skip > 0:
                            st.info(f"ℹ️ Đã bỏ qua {count_skip} file cũ (đã có trong bộ nhớ).")

        # 2. DUYỆT DRIVE (CÓ CHECK TRÙNG LẶP)
        elif source_option == "📁 Duyệt Google Drive":
            service = get_drive_service()
            if service:
                if not st.session_state.folder_tree_cache:
                    with st.spinner("Đang quét cấu trúc..."):
                        tree = [{'id': ROOT_FOLDER_ID, 'name': '📂 Thư mục gốc'}]
                        tree.extend(get_all_folders_recursive(service, ROOT_FOLDER_ID))
                        st.session_state.folder_tree_cache = tree
                
                folder_map = {item['name']: item['id'] for item in st.session_state.folder_tree_cache}
                selected_folder_name = st.selectbox("Chọn Chủ đề:", list(folder_map.keys()))
                selected_folder_id = folder_map[selected_folder_name]
                
                st.session_state.current_folder_id = selected_folder_id
                st.session_state.current_folder_name = selected_folder_name

                read_mode = st.radio("Phạm vi:", ["File trong thư mục này", "🚀 Quét sâu (Cả thư mục con)"])
                
                if st.button("🔍 Tìm file"):
                    with st.spinner("Đang tìm..."):
                        if read_mode == "🚀 Quét sâu (Cả thư mục con)":
                            files = list_files_deep(service, selected_folder_id)
                        else:
                            files = service.files().list(
                                q=f"'{selected_folder_id}' in parents and mimeType != 'application/vnd.google-apps.folder' and trashed=false",
                                fields="files(id, name, mimeType)", orderBy="name").execute().get('files', [])
                        st.session_state.found_files_cache = files
                        st.rerun()

                if st.session_state.found_files_cache:
                    files = st.session_state.found_files_cache
                    st.write(f"📂 Tìm thấy **{len(files)} file**.")
                    
                    max_val = len(files)
                    if max_val > 0:
                        limit = st.slider("Số lượng đọc:", 1, max_val, min(5, max_val))
                        
                        if st.button(f"📚 Đọc {limit} file"):
                            with st.spinner("Đang đọc và lọc dữ liệu cũ..."):
                                added_ctx = ""
                                count_new = 0
                                count_skip = 0
                                prog = st.progress(0)
                                
                                files_to_read = files[:limit]
                                for i, f in enumerate(files_to_read):
                                    # --- LOGIC CHỐNG TRÙNG ---
                                    if f['name'] in st.session_state.read_history:
                                        count_skip += 1
                                        prog.progress((i+1)/limit)
                                        continue

                                    try:
                                        content = read_drive_file(service, f['id'], f['name'], f['mimeType'])
                                        if len(content) > 50:
                                            added_ctx += f"\n=== TÀI LIỆU DRIVE: {f['name']} ===\n{content}\n"
                                            st.session_state.read_history.append(f['name'])
                                            count_new += 1
                                    except: pass
                                    prog.progress((i+1)/limit)
                                
                                st.session_state.global_context += added_ctx
                                
                                msg = ""
                                if count_new > 0:
                                    msg += f"✅ **Đã nạp thêm {count_new} tài liệu mới.**\n"
                                if count_skip > 0:
                                    msg += f"ℹ️ **Đã bỏ qua {count_skip} tài liệu cũ** (tránh trùng lặp).\n"
                                
                                msg += f"\nTổng bộ nhớ hiện tại: {len(st.session_state.read_history)} file."
                                st.session_state.messages.append({"role": "assistant", "content": msg})
                                st.rerun()
                    else: st.warning("Không có file nào.")

    # ==========================================
    # CẤU HÌNH AI
    # ==========================================
    full_system_instruction = TNT_MASTER_PROMPT
    if st.session_state.global_context:
        full_system_instruction += f"\n\n[USER PROVIDED CONTEXT]:\n{st.session_state.global_context}"

    if "messages" not in st.session_state: st.session_state.messages = []

    # --- GIAO DIỆN CHAT ---
    st.title("🎓 TNT Dissertation Master AI")
    st.caption(f"📂 Vị trí: {st.session_state.current_folder_name} | 🧠 Đã nhớ: {len(st.session_state.read_history)} file")
    st.markdown("---")

    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]): st.markdown(msg["content"])

    # XỬ LÝ INPUT
    prompt = None
    if audio_bytes:
        with st.spinner("🎧 Đang nghe..."):
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".webm") as tf:
                    tf.write(audio_bytes['bytes']); tf_path = tf.name
                wav = tf_path.replace(".webm", ".wav")
                AudioSegment.from_file(tf_path).export(wav, format="wav")
                r = sr.Recognizer()
                with sr.AudioFile(wav) as s: prompt = r.recognize_google(r.record(s), language="vi-VN")
                os.remove(tf_path); os.remove(wav)
            except: st.warning("Lỗi Mic.")

    if not prompt: prompt = st.chat_input("Nhập yêu cầu...")

    if prompt:
        if not api_key: st.error("Thiếu API Key!"); st.stop()
        genai.configure(api_key=api_key)
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"): st.markdown(prompt)
        
        with st.chat_message("assistant"):
            ph = st.empty(); full_res = ""
            try:
                model = genai.GenerativeModel("models/gemini-2.0-flash", system_instruction=full_system_instruction)
                chat = model.start_chat(history=[{"role": m["role"], "parts": [m["content"]]} for m in st.session_state.messages if m["role"] != "system"])
                for chunk in chat.send_message(prompt, stream=True):
                    if chunk.text: full_res += chunk.text; ph.markdown(full_res + "▌")
                ph.markdown(full_res)
                st.session_state.messages.append({"role": "assistant", "content": full_res})
            except Exception as e: st.error(f"Lỗi AI: {e}")

    # TOOLS (CỐ ĐỊNH)
    if st.session_state.messages and st.session_state.messages[-1]["role"] == "assistant":
        last_msg = st.session_state.messages[-1]["content"]
        st.divider()
        st.write("### 🛠️ Công cụ xử lý:")
        
        doc = Document(); doc.add_paragraph(last_msg); bio = BytesIO(); doc.save(bio); bio.seek(0)
        
        c1, c2, c3 = st.columns(3)
        with c1: st.download_button("📥 Tải về", data=bio, file_name="TNT_Output.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with c2:
            if st.button("☁️ Lưu vào Thư mục này"):
                with st.spinner("Lưu..."):
                    fid, fname = upload_to_drive(bio, "TNT_Output.docx", st.session_state.current_folder_id)
                    if fid: st.success(f"✅ Đã lưu vào '{st.session_state.current_folder_name}'!")
                    else: st.error(f"Lỗi: {fid}")
        with c3:
            if st.button("🔊 Đọc"):
                try:
                    tts = gTTS(text=last_msg, lang='vi'); mp3 = BytesIO(); tts.write_to_fp(mp3); st.audio(mp3, format='audio/mp3')
                except: pass

except Exception as e:
    st.error("🚨 HỆ THỐNG GẶP LỖI! Chi tiết:")
    st.code(traceback.format_exc())